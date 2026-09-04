from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "后端系统总结与优化报告.docx"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "EAF2F8"
PALE_BLUE = "F4F8FB"
GRAY = "667085"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D0D5DD"
TEXT = "1F2937"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color=MID_GRAY, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, fld_text, fld_end):
        run._r.append(el)


def set_east_asia_font(style, font_name="Microsoft YaHei") -> None:
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, size=None, bold=None, color=None, east_asia="Microsoft YaHei") -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_east_asia_font(normal)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.1

    specs = {
        "Title": (28, NAVY, 0, 12),
        "Subtitle": (14, GRAY, 0, 8),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        set_east_asia_font(style)
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("Caption", "Intense Quote"):
        style = doc.styles[name]
        set_east_asia_font(style)

    caption = doc.styles["Caption"]
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.font.italic = False
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)

    if "Body Lead" not in doc.styles:
        lead = doc.styles.add_style("Body Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = doc.styles["Body Lead"]
    set_east_asia_font(lead)
    lead.font.size = Pt(12)
    lead.font.color.rgb = RGBColor.from_string(NAVY)
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.2

    if "Callout" not in doc.styles:
        callout = doc.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = doc.styles["Callout"]
    set_east_asia_font(callout)
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = RGBColor.from_string(NAVY)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(7)
    callout.paragraph_format.space_after = Pt(9)
    callout.paragraph_format.line_spacing = 1.15


def add_bullet_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    existing_abs = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs, default=-1) + 1
    next_num = max(existing_num, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(next_abs))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Segoe UI Symbol")
    r_fonts.set(qn("w:hAnsi"), "Segoe UI Symbol")
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    first_num_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(next_abs))
    num.append(ref)
    numbering.append(num)
    return next_num


def add_list_item(doc: Document, text: str, num_id: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    p.add_run(text)


def add_labeled_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(text)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(table, color=PALE_BLUE, size="0")
    table.columns[0].width = Inches(0.08)
    table.columns[1].width = Inches(6.42)
    left, body = table.rows[0].cells
    set_cell_width(left, 115)
    set_cell_width(body, 9245)
    set_cell_shading(left, BLUE)
    set_cell_shading(body, PALE_BLUE)
    set_cell_margins(left, 80, 0, 80, 0)
    set_cell_margins(body, 120, 180, 120, 180)
    p = body.paragraphs[0]
    p.style = doc.styles["Callout"]
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run(text)
    prevent_row_split(table.rows[0])


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    caption: str | None = None,
    page_break_before: bool = False,
    compact: bool = False,
) -> None:
    if caption:
        p = doc.add_paragraph(caption, style="Caption")
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.page_break_before = page_break_before
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(table)
    table._tbl.tblPr.append(OxmlElement("w:tblLayout"))
    table._tbl.tblPr[-1].set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(OxmlElement("w:tblW"))
    table._tbl.tblPr[-1].set(qn("w:w"), "9360")
    table._tbl.tblPr[-1].set(qn("w:type"), "dxa")
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell, top=60 if compact else 80, bottom=60 if compact else 80)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if compact:
            p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
        if compact:
            set_run_font(run, size=9.5, bold=True, color=NAVY)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell, top=60 if compact else 80, bottom=60 if compact else 80)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if compact:
                p.paragraph_format.line_spacing = 1.0
            run = p.add_run(text)
            if compact:
                set_run_font(run, size=9.5, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_sections(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    cover_section = doc.sections[0]
    configure_page(cover_section)
    cover_section.different_first_page_header_footer = False


def start_body_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(section)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), "1")
    section._sectPr.append(pg_num_type)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)
    for run in fp.runs:
        set_run_font(run, size=8.5, color=GRAY)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(40)
    r = p.add_run("PROJECT REVIEW  ·  2026")
    set_run_font(r, size=9, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("后端系统总结\n与优化报告")
    set_run_font(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Advertiser CRM Backend")
    set_run_font(r, size=15, color=GRAY)

    divider = doc.add_table(rows=1, cols=1)
    divider.autofit = False
    divider.columns[0].width = Inches(6.5)
    set_cell_width(divider.cell(0, 0), 9360)
    set_cell_shading(divider.cell(0, 0), BLUE)
    set_cell_margins(divider.cell(0, 0), 20, 0, 20, 0)
    set_table_borders(divider, color=BLUE, size="0")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("报告范围")
    set_run_font(r, size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("Sprint 1 — Sprint 3 完整开发过程")
    set_run_font(r, size=14, bold=True, color=NAVY)

    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    widths = [1800, 7560]
    items = [
        ("技术基线", "Java 21  ·  Spring Boot 3.5  ·  PostgreSQL 16  ·  MyBatis-Plus  ·  Flyway"),
        ("报告主题", "系统架构  ·  技术难点  ·  优化说明  ·  项目亮点  ·  未来扩展"),
        ("形成时间", "2026 年 9 月"),
    ]
    set_table_borders(table, color=WHITE, size="0")
    for row, (label, value) in zip(table.rows, items):
        prevent_row_split(row)
        for i, text in enumerate((label, value)):
            cell = row.cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell, 100, 0, 100, 120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_run_font(run, size=9.5, bold=(i == 0), color=BLUE if i == 0 else TEXT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("本报告依据当前代码、数据库迁移、性能复现实验与自动化测试结果整理。")
    set_run_font(r, size=9, color=GRAY)


def add_architecture(doc: Document, bullets: int) -> None:
    heading = doc.add_heading("1  系统架构总结", level=1)
    heading.paragraph_format.page_break_before = True
    p = doc.add_paragraph(style="Body Lead")
    p.add_run(
        "Advertiser CRM 是面向广告业务后台的模块化单体系统。项目围绕内部用户、广告主档案、投放数据、统计报表、账户资金与充值处理形成完整业务闭环；三个 Sprint 依次完成基础能力、核心业务和工程化收尾。"
    )

    doc.add_heading("1.1  三个 Sprint 的架构演进", level=2)
    add_table(
        doc,
        ["阶段", "建设重点", "架构结果"],
        [
            ["Sprint 1\n基础建设", "公共 Web 规范、JWT 认证与权限、用户管理、广告主分类和档案", "建立模块化分包、统一响应、角色控制、Flyway 基线与数据库约束，为后续业务模块提供稳定入口。"],
            ["Sprint 2\n业务闭环", "投放数据、统计报表、广告主账户、消费流水、充值订单和模拟支付回调", "从普通 CRUD 扩展到聚合 SQL、资金事务、并发控制、状态机、HMAC 验签与幂等处理。"],
            ["Sprint 3\n工程收尾", "统一错误语义、参数校验、全局异常、性能优化、日志、结构与文档整理", "在不新增业务范围的前提下，提高一致性、可观测性、可维护性与可复现性。"],
        ],
        [1200, 3180, 4980],
        "表 1  Sprint 1—Sprint 3 的开发主线",
    )

    doc.add_heading("1.2  分层与模块边界", level=2)
    doc.add_paragraph(
        "系统采用一个 Spring Boot 应用与一个 PostgreSQL 数据库。代码按业务模块组织，模块内部再按 Controller、Service、Mapper、DTO、Entity、Exception、Validation 等职责分层。Controller 负责协议适配与输入校验，Service 负责业务规则和事务，Mapper/MyBatis XML 负责持久化与聚合查询，数据库约束负责兜底数据完整性。"
    )
    for text in (
        "公共层：common、config、security，统一响应、错误码、异常处理、请求追踪与安全配置。",
        "基础管理：auth、user、category、advertiser，负责身份、权限和广告主主数据。",
        "业务事实：delivery、report，负责投放记录、筛选分页和统计聚合。",
        "资金链路：account、payment，负责余额、不可变流水、充值订单与回调审计。",
    ):
        add_list_item(doc, text, bullets)

    doc.add_heading("1.3  业务与数据流", level=2)
    doc.add_paragraph(
        "广告主是系统的业务聚合入口：广告主连接分类与负责人，拥有投放事实和唯一资金账户；投放事实进入报表聚合，也可以被消费流水引用；资金账户连接消费流水、充值订单和支付回调审计。跨请求追踪与幂等依赖外部记录号、业务号、订单号和支付事件号，数据库自增 ID 仅用于内部关联。"
    )
    add_labeled_paragraph(doc, "管理链路：", "用户认证与授权 → 分类/广告主建档 → 负责人及状态维护。")
    add_labeled_paragraph(doc, "投放链路：", "投放事实录入 → 组合筛选与物理分页 → 日/周/月及多维报表聚合。")
    add_labeled_paragraph(doc, "资金链路：", "创建零余额账户 → 原子消费扣款/充值入账 → 追加不可变流水 → 审计追踪。")

    doc.add_heading("1.4  数据完整性与事务边界", level=2)
    doc.add_paragraph(
        "Flyway 以只增不改的迁移方式完成数据库从 V1 到 V11 的演进。应用校验用于尽早返回清晰错误，数据库的非空、唯一、检查约束、外键与触发器则抵御并发请求、代码遗漏或绕过应用的非法写入。"
    )
    for text in (
        "创建广告主与零余额账户在同一事务完成。",
        "余额扣减与消费流水追加在同一事务完成，失败时整体回滚。",
        "支付成功同时更新订单、增加余额、写充值流水和回调审计。",
        "管理员状态变更锁定有效管理员集合，保证系统始终保留可用管理员。",
    ):
        add_list_item(doc, text, bullets)
    add_callout(doc, "架构取舍：当前规模优先采用模块化单体，保持部署与本地事务简单；通过清晰模块边界和数据库保护获得可维护性，而不是在业务规则尚未稳定时过早拆分微服务。")


def add_challenges(doc: Document, bullets: int) -> None:
    heading = doc.add_heading("2  核心技术难点分析", level=1)
    heading.paragraph_format.page_break_before = True
    p = doc.add_paragraph(style="Body Lead")
    p.add_run("技术难点主要集中在权限的实时有效性、统计口径、资金一致性以及公开支付回调的安全与幂等。它们贯穿 Sprint 1 和 Sprint 2，并在 Sprint 3 通过契约测试与工程规范进一步固化。")

    challenges = [
        (
            "2.1  认证、权限与账号状态",
            "难点",
            "JWT 签名有效并不代表账号当前仍可用；用户可能在 Token 有效期内被禁用或改变角色。管理员变更还必须避免最后一个有效管理员被并发禁用或删除。",
            [
                "认证过滤器解析 Token 后重新读取数据库，仅为 ACTIVE 用户建立安全上下文，角色以数据库当前状态为准。",
                "公开注册固定创建 PENDING OPERATOR，客户端不能自行申请管理员权限。",
                "降级、禁用或删除管理员前锁定有效管理员集合，并用数据库保护作为最终防线。",
                "登录与注册入口限流，登录失败使用统一凭据错误，避免泄露账号是否存在。",
            ],
            "结果：权限判断可反映账号实时状态，管理员入口在并发场景下仍可保持可用。",
        ),
        (
            "2.2  统计口径与动态聚合查询",
            "难点",
            "CTR、CVR、CPC 等指标不能先按明细计算再平均；查询同时存在可选日期、广告主、广告类型、时间粒度、排序和分页，容易出现口径漂移、分页总数不一致或不安全动态 SQL。",
            [
                "在 PostgreSQL 中先聚合分子和分母，再计算比率；通过 NULLIF 与 COALESCE 稳定处理零分母。",
                "日期范围、粒度、排序字段和方向在 Service 层归一化并使用枚举白名单。",
                "数据查询与 COUNT 复用相同过滤条件，趋势按自然日、周一为起点的自然周或自然月聚合。",
                "仅在存在广告类型筛选时 JOIN 类型表，兼顾正确性和执行效率。",
            ],
            "结果：总览、趋势和维度报表具有统一统计口径，支持稳定物理分页与受控排序。",
        ),
        (
            "2.3  账户余额、流水与并发扣款",
            "难点",
            "“先查余额再扣款”会在并发下超扣；余额与流水若分开提交会产生账实不一致；重复业务号还可能造成重复扣款。",
            [
                "使用 UPDATE … WHERE balance >= amount RETURNING balance 原子判断并扣减余额。",
                "扣款和 CONSUMPTION 流水写入位于同一事务，后续失败自动回滚。",
                "business_no 全局唯一，并使用条件写入处理并发重复请求。",
                "资金流水只追加、不开放修改和删除；关联投放时锁定记录并校验广告主归属。",
            ],
            "结果：并发扣款不会突破余额，重复请求不会形成重复资金变动，余额与流水可相互核对。",
        ),
        (
            "2.4  支付回调验签、状态机与幂等",
            "难点",
            "支付回调是无 JWT 的公开入口，需要同时抵御伪造、篡改、重放、重复或并发事件，并保证订单、余额、流水和审计记录的一致性。",
            [
                "使用 HMAC-SHA256(secret, timestamp + '.' + rawBody) 对原始字节验签，并限制时间窗口与请求体大小。",
                "订单只允许 PENDING 向 SUCCESS、FAILED 或 CLOSED 迁移，终态不可回退。",
                "provider_event_id、business_no、recharge_order_id 和平台交易号形成多层唯一约束。",
                "成功处理在一个事务内完成订单更新、余额增加、充值流水和回调审计；重复事件返回幂等结果，载荷冲突明确拒绝。",
            ],
            "结果：模拟支付虽不连接真实渠道，但已验证真实支付接入所需的核心安全边界与一致性模型。",
        ),
    ]
    for index, (title, label, problem, solutions, result) in enumerate(challenges):
        heading = doc.add_heading(title, level=2)
        if index == 1:
            heading.paragraph_format.page_break_before = True
        add_labeled_paragraph(doc, f"{label}：", problem)
        add_labeled_paragraph(doc, "解决思路：", "")
        for item in solutions:
            add_list_item(doc, item, bullets)
        add_callout(doc, result)


def add_optimizations(doc: Document, bullets: int) -> None:
    heading = doc.add_heading("3  优化点说明", level=1)
    heading.paragraph_format.page_break_before = True
    p = doc.add_paragraph(style="Body Lead")
    p.add_run("Sprint 3 的目标不是继续增加业务功能，而是对前两个 Sprint 的实现进行工程化完善。优化覆盖接口契约、非法输入防护、异常兜底、核心 SQL、日志和代码结构，并以自动化测试与可复现实验确认行为未发生回退。")

    doc.add_heading("3.1  统一返回、错误码与全局异常", level=2)
    for text in (
        "业务接口统一返回 ApiResponse<T> 或 ResponseEntity<ApiResponse<T>>；分页数据使用 PageResponse<T>。",
        "响应固定包含 success、code、message、data、timestamp 和 requestId，便于客户端处理与问题追踪。",
        "公共错误码使用 COMMON_ 前缀，业务错误码使用所属模块前缀；契约测试检查全局唯一性、命名格式和 HTTP 状态映射。",
        "全局异常处理覆盖 Bean Validation、参数绑定、类型转换、非法 JSON、方法/媒体类型错误、数据库冲突和未知异常；401/403 保持相同 JSON 契约。",
    ):
        add_list_item(doc, text, bullets)

    doc.add_heading("3.2  参数校验与非法输入防护", level=2)
    doc.add_paragraph(
        "校验按 DTO、Service、数据库三层布置：DTO 负责必填、长度、格式、范围和金额精度；Service 负责跨字段、状态迁移与跨实体关系；数据库负责并发和绕过应用场景。局部更新先合并旧值再验证完整业务关系，避免只验证本次提交字段而留下非法状态。"
    )

    doc.add_heading("3.3  核心接口性能优化", level=2)
    doc.add_paragraph(
        "性能优化使用 PostgreSQL 16、500,000 条固定投放记录和相同查询参数，通过 EXPLAIN (ANALYZE, BUFFERS) 比较前后执行计划。优化只针对已确认的查询模式，不引入缓存，也不改变接口结果。"
    )
    add_table(
        doc,
        ["测试场景", "优化前", "优化后", "改善"],
        [
            ["投放组合筛选分页", "0.418 ms\n464 个共享缓冲页", "0.191 ms\n27 个共享缓冲页", "时间降低 54.3%\n缓冲页降低 94.2%"],
            ["31 天趋势聚合", "14.881 ms\n存在 Hash Join", "12.061 ms\n移除无用 JOIN", "时间降低 19.0%\n执行计划更直接"],
        ],
        [2580, 2040, 2040, 2700],
        "表 2  核心查询优化前后对比",
        page_break_before=True,
    )
    for text in (
        "V11 新增 (advertiser_id, advertising_type_id, record_date DESC, id DESC) 组合索引，使高频组合筛选与稳定排序直接匹配索引。",
        "报表 SQL 在没有广告类型条件时不再 JOIN advertising_types，减少无效读取与 Hash Join。",
        "性能脚本在独立数据库执行，并用事务回滚测试数据和临时对象，保证实验可重复且不污染开发数据。",
    ):
        add_list_item(doc, text, bullets)

    doc.add_heading("3.4  日志体系与可维护性", level=2)
    for text in (
        "请求过滤器生成或接收经过长度与字符限制的 Request ID，记录方法、路径、状态和耗时。",
        "ERROR 用于未知异常与 5xx，WARN 用于权限拒绝、冲突与限流，INFO 用于请求完成和关键业务状态，常规参数错误使用 DEBUG。",
        "日志不记录请求体、查询参数、Authorization、密码、JWT、密钥、签名或原始支付载荷。",
        "支付订单号与平台交易号的格式、长度和规范化规则集中到 PaymentReferenceRules，减少 Controller 与 Service 的重复逻辑。",
        "保持按业务模块分包，复杂报表 SQL 放入 MyBatis XML，事务边界集中在 Service，使职责与变更影响范围更清晰。",
    ):
        add_list_item(doc, text, bullets)
    add_callout(doc, "验证结果：合并后的系统完成 420 项自动化回归，并通过容器启动、健康检查、OpenAPI、Flyway V11 与非 root 运行验证。")


def add_highlights(doc: Document) -> None:
    heading = doc.add_heading("4  项目亮点提炼", level=1)
    heading.paragraph_format.page_break_before = True
    p = doc.add_paragraph(style="Body Lead")
    p.add_run("项目亮点不在接口数量，而在于将权限、统计、资金和回调等高风险规则落实到可执行的事务、SQL、数据库约束和自动化测试中。")
    add_table(
        doc,
        ["项目亮点", "具体体现", "价值"],
        [
            ["业务规则分层保护", "DTO 提前反馈、Service 统一业务规则、PostgreSQL 约束与触发器兜底", "既改善接口体验，也能抵御并发、遗漏和绕过应用的非法写入。"],
            ["资金链路可追踪", "原子余额更新、不可变流水、唯一业务号、事务回滚与关联一致性", "防止超扣和重复扣款，并可根据流水核对每次余额变化。"],
            ["支付回调安全模型完整", "原始字节 HMAC、时间窗口、载荷限制、状态机、事件幂等和拒绝审计", "为未来接入真实支付渠道保留清晰且可验证的安全边界。"],
            ["统计口径稳定", "数据库聚合、零分母处理、白名单排序、物理分页和相同过滤条件 COUNT", "避免指标口径漂移，确保总览、趋势和维度报表结果一致。"],
            ["权限反映实时状态", "JWT 验签后重新读取账号状态与角色，最后一个管理员受并发锁与数据库保护", "Token 未过期也不能绕过账号禁用，系统不会丢失管理入口。"],
            ["优化有证据可复现", "50 万条固定数据、前后执行计划、组合索引和可回滚实验脚本", "优化结论来自测量，不依赖主观判断，也不以缓存掩盖查询问题。"],
            ["工程规范可防回归", "错误码、响应类型、HTTP 状态、日志等级与敏感信息均有自动化检查", "把代码规范转化为持续可执行的质量门槛。"],
            ["项目可独立验证", "Maven Wrapper、Flyway、Testcontainers、Dockerfile、Compose 与 420 项回归", "降低环境差异，提高交接、复现和后续扩展效率。"],
        ],
        [2200, 3980, 3180],
        "表 3  项目亮点与实际价值",
        compact=True,
    )
    add_callout(doc, "总体特点：先保证业务正确性和数据一致性，再依据可复现证据进行性能与工程优化；技术选择围绕真实问题，不为展示复杂度而引入不必要组件。")


def add_future(doc: Document, bullets: int) -> None:
    heading = doc.add_heading("5  未来扩展方向", level=1)
    heading.paragraph_format.page_break_before = True
    p = doc.add_paragraph(style="Body Lead")
    p.add_run("现有实现满足本次后端实习项目范围。后续扩展应沿用当前模块边界和一致性原则，优先补齐真实运行条件，再根据数据量、团队与负载证据决定是否引入分布式组件。")

    directions = [
        (
            "5.1  认证、安全与审计",
            [
                "将进程内认证限流迁移到 Redis 等共享存储，并在可信反向代理边界统一解析客户端地址。",
                "引入 Refresh Token 轮换、Token ID 撤销与签名密钥版本管理。",
                "增加管理操作审计，记录操作人、对象、变更前后摘要和 Request ID；对登录失败、权限拒绝和支付异常建立指标与告警。",
            ],
        ),
        (
            "5.2  真实支付与资金运营",
            [
                "在现有订单、回调接口和状态机后增加支付渠道适配器，使账户模块不依赖具体渠道 SDK。",
                "增加主动查单、退款、日终对账和差错处理，所有外部动作继续使用业务号与幂等键。",
                "跨系统可靠投递可采用 Outbox，再由异步任务发送，避免数据库事务与消息发送双写不一致。",
            ],
        ),
        (
            "5.3  投放数据接入与报表规模化",
            [
                "增加批量导入、文件校验、错误行报告、可重试任务和数据质量规则，以承接渠道级数据。",
                "数据增长后优先按业务日期分区，并依据慢查询证据选择预聚合表或物化视图。",
                "仅当分析负载明显影响交易库时，再通过 CDC 或批处理同步到独立分析存储。",
            ],
        ),
        (
            "5.4  工程交付与可观测性",
            [
                "建立 CI 流水线，自动执行编译、完整回归、迁移测试、镜像构建和依赖安全检查。",
                "共享环境补充 TLS、外部 Secret、最小权限数据库账号、备份恢复演练和健康告警。",
                "接入结构化日志、指标与链路追踪，以 Request ID 关联接口、SQL 和外部调用。",
                "继续保持模块化单体；只有团队、部署频率或负载边界明确分离时，再评估拆分报表、数据接入或支付适配模块。",
            ],
        ),
    ]
    for index, (title, items) in enumerate(directions):
        heading = doc.add_heading(title, level=2)
        if index == 3:
            heading.paragraph_format.page_break_before = True
        for item in items:
            add_list_item(doc, item, bullets)

    add_callout(doc, "最终结论：三个 Sprint 已完成从管理基础、业务闭环到工程化完善的完整开发过程。系统当前具备清晰架构、可验证业务规则、可追踪资金链路、可复现性能优化与持续扩展基础。")


def build() -> None:
    doc = Document()
    configure_sections(doc)
    configure_styles(doc)
    bullets = add_bullet_numbering(doc)
    props = doc.core_properties
    props.title = "Advertiser CRM 后端系统总结与优化报告"
    props.subject = "Sprint 1 至 Sprint 3 系统架构、技术难点、优化、亮点与扩展方向"
    props.author = "Advertiser CRM Project"
    props.keywords = "Advertiser CRM, Spring Boot, Sprint, 系统总结, 优化报告"
    props.comments = "Final DOCX deliverable"

    add_cover(doc)
    start_body_section(doc)
    add_architecture(doc, bullets)
    add_challenges(doc, bullets)
    add_optimizations(doc, bullets)
    add_highlights(doc)
    add_future(doc, bullets)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
